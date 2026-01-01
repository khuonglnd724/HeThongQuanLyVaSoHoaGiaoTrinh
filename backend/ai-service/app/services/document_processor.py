"""Document processing for PDF and Word documents"""
import logging
from typing import Dict, List, Optional
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process PDF and Word documents for text extraction"""
    
    @staticmethod
    def extract_from_pdf(file_path: str, max_pages: Optional[int] = None) -> Dict[str, str]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            max_pages: Maximum number of pages to extract (None = all)
            
        Returns:
            Dict with 'full_text', 'pages', 'metadata'
        """
        try:
            full_text = []
            pages_content = []
            
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                metadata = pdf.metadata or {}
                
                # Extract text from pages
                for idx, page in enumerate(pdf.pages):
                    if max_pages and idx >= max_pages:
                        break
                    
                    page_text = page.extract_text()
                    if page_text:
                        pages_content.append({
                            "page_number": idx + 1,
                            "content": page_text.strip()
                        })
                        full_text.append(page_text)
                
                return {
                    "full_text": "\n\n".join(full_text),
                    "pages": pages_content,
                    "metadata": {
                        "title": metadata.get("Title", ""),
                        "author": metadata.get("Author", ""),
                        "pages": len(pdf.pages),
                        "file_type": "pdf"
                    }
                }
        
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, str]:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with 'full_text', 'sections', 'metadata'
        """
        try:
            doc = DocxDocument(file_path)
            
            full_text = []
            sections = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
                    sections.append({
                        "type": "paragraph",
                        "level": para.style.name,
                        "content": text
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_data))
                
                table_content = "\n".join(table_text)
                full_text.append(f"[Bảng {table_idx + 1}]\n{table_content}")
                sections.append({
                    "type": "table",
                    "index": table_idx + 1,
                    "content": table_content
                })
            
            return {
                "full_text": "\n\n".join(full_text),
                "sections": sections,
                "metadata": {
                    "file_type": "docx",
                    "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                    "tables": len(doc.tables)
                }
            }
        
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_document(file_path: str, max_pages: Optional[int] = None) -> Dict[str, str]:
        """
        Auto-detect document type and extract text
        
        Args:
            file_path: Path to document file
            max_pages: Max pages for PDF (None = all)
            
        Returns:
            Dict with extracted content and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_type = path.suffix.lower()
        
        try:
            if file_type == ".pdf":
                return DocumentProcessor.extract_from_pdf(file_path, max_pages)
            elif file_type in [".docx", ".doc"]:
                return DocumentProcessor.extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove multiple spaces
        text = " ".join(text.split())
        
        # Remove special characters but keep punctuation
        import re
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 1000,
        overlap: int = 100
    ) -> List[Dict[str, str]]:
        """
        Split text into chunks for embedding
        
        Args:
            text: Full text to chunk
            chunk_size: Characters per chunk
            overlap: Character overlap between chunks
            
        Returns:
            List of dicts with 'content' and 'index'
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            
            # Try to break at sentence boundary
            if end < len(text):
                # Find last period before chunk_size
                last_period = text.rfind('. ', start, end)
                if last_period > start + chunk_size // 2:
                    end = last_period + 2
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "content": chunk_text,
                    "index": len(chunks),
                    "start": start,
                    "end": end
                })
            
            # Move start position (with overlap)
            start = end - overlap
        
        return chunks
